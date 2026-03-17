import asyncio
from typing import Optional
import subprocess
import os
import platform

# Text extraction imports
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def _extract_text_with_ocr(file_path: str) -> str:
    """Fallback OCR extraction for scanned/image-based PDFs."""
    try:
        import pytesseract  # type: ignore
        import pypdfium2 as pdfium  # type: ignore
        from PIL import Image  # type: ignore
    except ImportError as err:
        print(f"[PDF OCR] OCR dependencies missing ({err}). Install pytesseract and pypdfium2.")
        return ""
    
    try:
        pdf = pdfium.PdfDocument(file_path)
        text_chunks = []
        for page_index, page in enumerate(pdf):
            try:
                # Render page to high-res bitmap for better OCR accuracy
                bitmap = page.render(scale=2.0, rotation=0)
                pil_image: Image.Image = bitmap.to_pil()
                ocr_text = pytesseract.image_to_string(pil_image)
                if ocr_text.strip():
                    text_chunks.append(ocr_text)
                else:
                    print(f"[PDF OCR] Warning: OCR returned empty text for page {page_index + 1}")
            except Exception as page_error:
                print(f"[PDF OCR] Failed to OCR page {page_index + 1}: {page_error}")
        combined_text = "\n".join(text_chunks).strip()
        if not combined_text:
            print(f"[PDF OCR] Warning: OCR produced no text for {file_path}")
        else:
            print(f"[PDF OCR] Successfully extracted {len(combined_text)} characters via OCR fallback")
        return combined_text
    except Exception as ocr_error:
        print(f"[PDF OCR] Unexpected error: {ocr_error}")
        return ""


async def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF using pdfplumber, with OCR fallback for image-only PDFs."""
    if not PDF_AVAILABLE:
        raise ImportError("pdfplumber not available. Install with: pip install pdfplumber")
    
    def _extract():
        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    else:
                        print(f"[PDF EXTRACT] Warning: Page {page_num + 1} extracted no text")
            
            if not text.strip():
                print(f"[PDF EXTRACT] No text extracted; trying OCR for {file_path}")
                text = _extract_text_with_ocr(file_path)
            
            if text.strip():
                print(f"[PDF EXTRACT] Successfully extracted {len(text)} characters from PDF")
            return text
        except FileNotFoundError:
            raise Exception(f"PDF file not found: {file_path}")
        except Exception as e:
            # Try OCR before failing
            print(f"[PDF EXTRACT] Error: {e}; trying OCR fallback")
            ocr_text = _extract_text_with_ocr(file_path)
            return ocr_text
    
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _extract)


async def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not available")
    
    def _extract():
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_texts.append(" | ".join(row_text))
            
            # Combine all text
            all_text = paragraphs + table_texts
            extracted_text = '\n'.join(all_text)
            
            # If still no text, try to extract from runs and other elements
            if not extracted_text.strip():
                fallback_text = []
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        run_text = run.text.strip()
                        if run_text:
                            fallback_text.append(run_text)
                extracted_text = ' '.join(fallback_text)
            
            return extracted_text
            
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    # Run in thread pool to avoid blocking
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _extract)


async def extract_text_from_doc(file_path: str) -> str:
    """Extract text from DOC file using antiword"""
    def _extract():
        try:
            if platform.system() == "Windows":
                # Try antiword if available
                result = subprocess.run(
                    ['antiword', file_path],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode == 0:
                    return result.stdout
            else:
                # Linux/Mac antiword
                result = subprocess.run(
                    ['antiword', file_path],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode == 0:
                    return result.stdout
            
            # Fallback: try to read as plain text
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
                
        except Exception:
            # Last resort: try binary read and decode
            try:
                with open(file_path, 'rb') as f:
                    content = f.read()
                    return content.decode('utf-8', errors='ignore')
            except Exception as e:
                raise Exception(f"Failed to extract text from DOC: {str(e)}")
    
    # Run in thread pool to avoid blocking
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _extract)


async def extract_text_from_txt(file_path: str) -> str:
    """Extract text from TXT file"""
    def _extract():
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
        except Exception as e:
            raise Exception(f"Failed to extract text from TXT: {str(e)}")
    
    # Run in thread pool to avoid blocking
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _extract)


async def extract_text_from_file(file_path: str) -> str:
    """Extract text from file based on extension"""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    file_ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_ext == '.pdf':
            text = await extract_text_from_pdf(file_path)
            # If still empty after OCR, raise to trigger plain-text fallback below
            if not text.strip():
                raise Exception("No text extracted from PDF (even after OCR)")
            return text
        elif file_ext == '.docx':
            return await extract_text_from_docx(file_path)
        elif file_ext == '.doc':
            return await extract_text_from_doc(file_path)
        elif file_ext == '.txt':
            return await extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
    
    except Exception as e:
        # Fallback: try to read as plain text
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
                if content.strip():
                    return content
        except Exception:
            pass
        
        raise Exception(f"Failed to extract text: {str(e)}")
