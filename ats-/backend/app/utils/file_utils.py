import os
import tempfile
import zipfile
from typing import List, Tuple
from fastapi import UploadFile, HTTPException, status
import platform
import asyncio

# Platform-specific magic import
try:
    if platform.system() == "Windows":
        import magic
        # For python-magic-bin on Windows
        magic_instance = magic.Magic(mime=True)
    else:
        import magic
        # For python-magic on Linux/Mac
        magic_instance = magic.Magic(mime=True)
except ImportError:
    magic_instance = None

ALLOWED_EXTENSIONS = {'.pdf', '.doc', '.docx', '.txt'}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB


def get_file_extension(filename: str) -> str:
    """Get file extension"""
    return os.path.splitext(filename)[1].lower()


def validate_resume_file(file: UploadFile) -> bool:
    """Validate uploaded resume file"""
    # Check file extension
    ext = get_file_extension(file.filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File type {ext} not allowed. Allowed types: {', '.join(ALLOWED_EXTENSIONS)}"
        )
    
    # Check file size
    file.file.seek(0, 2)  # Seek to end
    file_size = file.file.tell()
    file.file.seek(0)  # Reset to beginning
    
    if file_size > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail="File size too large. Maximum size is 10MB"
        )
    
    # Validate MIME type if magic is available
    if magic_instance:
        file_content = file.file.read(2048)  # Read first 2KB
        file.file.seek(0)  # Reset
        
        mime_type = magic_instance.from_buffer(file_content)
        valid_mimes = {
            'application/pdf',
            'application/msword',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'text/plain'
        }
        
        if mime_type not in valid_mimes:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Invalid file type detected: {mime_type}"
            )
    
    return True


async def save_upload_file_tmp(upload_file: UploadFile) -> str:
    """Save uploaded file to temporary location"""
    suffix = get_file_extension(upload_file.filename)
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
        content = await upload_file.read()
        tmp_file.write(content)
        tmp_file.flush()
        return tmp_file.name


def extract_zip_and_filter(zip_path: str) -> List[str]:
    """Extract ZIP file and return list of resume files"""
    resume_files = []
    extract_dir = tempfile.mkdtemp()
    
    try:
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            zip_ref.extractall(extract_dir)
            
        # Walk through extracted files
        for root, dirs, files in os.walk(extract_dir):
            for file in files:
                file_path = os.path.join(root, file)
                ext = get_file_extension(file)
                
                if ext in ALLOWED_EXTENSIONS:
                    resume_files.append(file_path)
                    
    except zipfile.BadZipFile:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid ZIP file"
        )
    
    return resume_files


async def extract_text_from_file(file_path: str) -> str:
    """Extract text content from various file formats"""
    try:
        file_extension = get_file_extension(file_path)
        
        if file_extension == '.pdf':
            return await extract_text_from_pdf(file_path)
        elif file_extension in ['.doc', '.docx']:
            return await extract_text_from_doc(file_path)
        elif file_extension == '.txt':
            return await extract_text_from_txt(file_path)
        else:
            return ""
    except Exception as e:
        print(f"Error extracting text from {file_path}: {str(e)}")
        return ""


async def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file using pdfplumber (better layout preservation)"""
    try:
        import pdfplumber
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # pdfplumber preserves layout better than PyPDF2 automatically
                # This prevents text concatenation issues like "LANGUAGES7026720645yogesh"
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                else:
                    print(f"[PDF EXTRACT] Warning: Page {page_num + 1} extracted no text")
        
        if not text.strip():
            print(f"[PDF EXTRACT] Warning: No text extracted from PDF: {file_path}")
            ocr_text = _extract_text_with_ocr(file_path)
            if ocr_text:
                print(f"[PDF OCR] Successfully extracted {len(ocr_text)} characters via OCR fallback")
            return ocr_text
        
        print(f"[PDF EXTRACT] Successfully extracted {len(text)} characters using pdfplumber")
        return text
    except ImportError:
        print("Error: pdfplumber not installed. Install with: pip install pdfplumber")
        return _extract_text_with_ocr(file_path)
    except FileNotFoundError:
        print(f"Error: PDF file not found: {file_path}")
        return ""
    except Exception as e:
        print(f"Error extracting PDF text: {str(e)}")
        import traceback
        traceback.print_exc()
        ocr_text = _extract_text_with_ocr(file_path)
        if ocr_text:
            print(f"[PDF OCR] Successfully extracted {len(ocr_text)} characters via OCR fallback after error")
        return ocr_text


def _extract_text_with_ocr(file_path: str) -> str:
    """Fallback OCR extraction for scanned/image-based PDFs."""
    try:
        import pytesseract
        import pypdfium2 as pdfium
        from PIL import Image
    except ImportError as err:
        print(f"[PDF OCR] OCR dependencies missing ({err}). Install pytesseract and pypdfium2.")
        return ""
    
    try:
        pdf = pdfium.PdfDocument(file_path)
        text_chunks: List[str] = []
        for page_index, page in enumerate(pdf):
            try:
                # Render page to high-resolution bitmap for better OCR accuracy
                bitmap = page.render(scale=2.0, rotation=0)
                pil_image: Image.Image = bitmap.to_pil()
                ocr_text = pytesseract.image_to_string(pil_image)
                if ocr_text.strip():
                    text_chunks.append(ocr_text)
                else:
                    print(f"[PDF OCR] Warning: OCR returned empty text for page {page_index + 1}")
            except Exception as page_error:
                print(f"[PDF OCR] Failed to OCR page {page_index + 1}: {page_error}")
        combined_text = "\n".join(text_chunks).strip()
        if not combined_text:
            print(f"[PDF OCR] Warning: OCR produced no text for {file_path}")
        return combined_text
    except Exception as ocr_error:
        print(f"[PDF OCR] Unexpected error: {ocr_error}")
        return ""


async def extract_text_from_doc(file_path: str) -> str:
    """Extract text from DOC/DOCX file"""
    try:
        if file_path.endswith('.docx'):
            from docx import Document
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        else:
            # For .doc files, we might need python-docx2txt or similar
            import docx2txt
            return docx2txt.process(file_path)
    except Exception as e:
        print(f"Error extracting DOC text: {str(e)}")
        return ""


async def extract_text_from_txt(file_path: str) -> str:
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        print(f"Error extracting TXT text: {str(e)}")
        return ""
